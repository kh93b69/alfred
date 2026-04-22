import json
import os
import logging
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook

from bot.services.ai import client
from bot.services.knowledge import load_knowledge
from bot.services import trello
from bot.services import task_memory
from bot.config import CLAUDE_MODEL, TRELLO_LIST_INBOX, TRELLO_LIST_DOING, TRELLO_LIST_DONE

logger = logging.getLogger(__name__)

# Хранилище предложенных задач (ожидающих подтверждения)
# Ключ — ID сообщения в Telegram, значение — данные задачи
pending_tasks: dict[str, dict] = {}

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


async def propose_daily_tasks() -> list[dict]:
    """
    Генерирует предложения задач на день.
    Возвращает список задач (НЕ создаёт карточки — ждёт подтверждения).
    """
    knowledge = load_knowledge()
    if not knowledge:
        return []

    # Текущие задачи
    current_cards = trello.get_cards()
    current_tasks = ""
    for card in current_cards:
        list_name = "Входящие"
        if card["idList"] == TRELLO_LIST_DOING:
            list_name = "В работе"
        elif card["idList"] == TRELLO_LIST_DONE:
            list_name = "Готово"
        current_tasks += f"- [{list_name}] {card['name']}\n"

    today = datetime.now().strftime("%d.%m.%Y, %A")
    history = task_memory.get_history_summary()

    prompt = f"""Ты — Альфред, автономный AI-ассистент. Сегодня {today}.

БАЗА ЗНАНИЙ ВЛАДЕЛЬЦА:
{knowledge}

ТЕКУЩИЕ ЗАДАЧИ НА ДОСКЕ:
{current_tasks if current_tasks else "Доска пуста"}

ИСТОРИЯ ЗАДАЧ (НЕ повторяй!):
{history if history else "История пуста"}

КРИТИЧЕСКИ ВАЖНО:
- НЕ предлагай задачи из истории "ОТКЛОНЁННЫЕ" — владелец их не хочет
- НЕ предлагай задачи из истории "ВЫПОЛНЕННЫЕ" — они уже сделаны
- НЕ повторяй задачи которые уже есть на доске
- Каждый день — СВЕЖИЕ, РАЗНЫЕ задачи
- Думай о том что НОВОЕ можно сделать сегодня для прогресса к целям
- Задачи должны быть разнообразными по типу: анализ, контент, стратегия, нетворкинг, рост

Сгенерируй 3-5 конкретных НОВЫХ задач на сегодня.

Для каждой задачи:
- name: краткое название (до 60 символов)
- description: что конкретно делать (2-3 предложения)
- output_type: "document" (docx), "table" (xlsx) или "text"

Верни ТОЛЬКО JSON массив:
[{{"name": "...", "description": "...", "output_type": "text"}}]"""

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}],
    )

    json_text = response.content[0].text.strip()
    if json_text.startswith("```"):
        json_text = json_text.split("\n", 1)[1]
        json_text = json_text.rsplit("```", 1)[0]

    tasks = json.loads(json_text)

    # Записываем что задачи были предложены
    for task in tasks:
        task_memory.add_proposed(task["name"])

    logger.info(f"Предложено {len(tasks)} задач на день")
    return tasks


def _generate_docx(title: str, content: str) -> str:
    """Создаёт .docx файл из текста"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Заголовок
    doc.add_heading(title, level=1)

    # Контент — разбиваем по строкам
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            doc.add_paragraph(line[2:].strip(), style="List Number")
        else:
            doc.add_paragraph(line)

    filepath = os.path.join(OUTPUT_DIR, f"{title[:50]}.docx")
    doc.save(filepath)
    return filepath


def _generate_xlsx(title: str, content: str) -> str:
    """Создаёт .xlsx файл — Claude генерирует JSON с данными таблицы"""
    wb = Workbook()
    ws = wb.active
    ws.title = title[:30]

    # Пытаемся распарсить как JSON таблицу
    try:
        # Ищем JSON в контенте
        json_start = content.find("[")
        json_end = content.rfind("]") + 1
        if json_start >= 0 and json_end > json_start:
            table_data = json.loads(content[json_start:json_end])
            if table_data and isinstance(table_data[0], dict):
                # Заголовки
                headers = list(table_data[0].keys())
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = cell.font.copy(bold=True)
                # Данные
                for row_idx, row_data in enumerate(table_data, 2):
                    for col, header in enumerate(headers, 1):
                        ws.cell(row=row_idx, column=col, value=row_data.get(header, ""))
            elif table_data and isinstance(table_data[0], list):
                for row_idx, row_data in enumerate(table_data, 1):
                    for col, value in enumerate(row_data, 1):
                        cell = ws.cell(row=row_idx, column=col, value=value)
                        if row_idx == 1:
                            cell.font = cell.font.copy(bold=True)
        else:
            raise ValueError("JSON не найден")
    except (json.JSONDecodeError, ValueError):
        # Если не JSON — просто кладём текст по строкам
        for row_idx, line in enumerate(content.split("\n"), 1):
            if line.strip():
                parts = [p.strip() for p in line.split("|") if p.strip()]
                if not parts:
                    parts = [p.strip() for p in line.split("\t") if p.strip()]
                if not parts:
                    parts = [line.strip()]
                for col, value in enumerate(parts, 1):
                    ws.cell(row=row_idx, column=col, value=value)

    # Автоширина колонок
    for col in ws.columns:
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

    filepath = os.path.join(OUTPUT_DIR, f"{title[:50]}.xlsx")
    wb.save(filepath)
    return filepath


async def execute_task(card_id: str, card_name: str, card_desc: str, output_type: str = "text") -> dict:
    """
    Выполняет задачу и генерирует результат нужного типа.
    Возвращает dict с результатом и путём к файлу (если есть).
    """
    knowledge = load_knowledge()

    # Формируем промпт в зависимости от типа результата
    if output_type == "table":
        format_instruction = (
            "Результат должен быть в виде таблицы. "
            "Верни данные как JSON массив объектов. Пример:\n"
            '[{"Колонка1": "значение", "Колонка2": "значение"}]\n'
            "Перед JSON можешь написать краткое описание (1-2 предложения)."
        )
    elif output_type == "document":
        format_instruction = (
            "Напиши полноценный документ с заголовками (используй # для заголовков), "
            "списками (используй - для пунктов) и структурированным текстом. "
            "Документ должен быть готов к использованию."
        )
    else:
        format_instruction = "Напиши конкретный, структурированный текстовый результат."

    prompt = f"""Ты — Альфред, автономный AI-ассистент. Выполни задачу.

БАЗА ЗНАНИЙ ВЛАДЕЛЬЦА:
{knowledge}

ЗАДАЧА: {card_name}
ОПИСАНИЕ: {card_desc}

{format_instruction}

Результат должен быть конкретным, полезным и готовым к использованию."""

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    content = response.content[0].text
    result = {"text": content, "file_path": None, "file_name": None}

    # Генерируем файл
    if output_type == "table":
        filepath = _generate_xlsx(card_name, content)
        result["file_path"] = filepath
        result["file_name"] = os.path.basename(filepath)
    elif output_type == "document":
        filepath = _generate_docx(card_name, content)
        result["file_path"] = filepath
        result["file_name"] = os.path.basename(filepath)

    # Перемещаем в "В работе"
    trello.move_to_doing(card_id)

    # Добавляем комментарий с кратким результатом
    preview = content[:500] + "..." if len(content) > 500 else content
    trello.add_comment(card_id, f"✅ Результат:\n\n{preview}")

    # Прикрепляем файл если есть
    if result["file_path"]:
        trello.attach_file(card_id, result["file_path"], result["file_name"])

    # Перемещаем в "На проверку"
    trello.move_to_review(card_id)

    # Запоминаем выполненную задачу
    task_memory.add_completed(card_name)

    logger.info(f"Задача выполнена: {card_name} (тип: {output_type})")
    return result


async def daily_report() -> str:
    """Формирует ежедневный отчёт"""
    cards = trello.get_cards()

    done = []
    review = []
    doing = []
    inbox = []

    for card in cards:
        if card["idList"] == TRELLO_LIST_DONE:
            done.append(card["name"])
        elif card["idList"] == trello.TRELLO_LIST_REVIEW:
            review.append(card["name"])
        elif card["idList"] == TRELLO_LIST_DOING:
            doing.append(card["name"])
        elif card["idList"] == TRELLO_LIST_INBOX:
            inbox.append(card["name"])

    report = f"📊 **Отчёт Альфреда**\n{datetime.now().strftime('%d.%m.%Y')}\n\n"

    if done:
        report += f"✅ Выполнено ({len(done)}):\n"
        for t in done:
            report += f"  • {t}\n"
        report += "\n"
    if review:
        report += f"👀 На проверку ({len(review)}):\n"
        for t in review:
            report += f"  • {t}\n"
        report += "\n"
    if doing:
        report += f"🔨 В работе ({len(doing)}):\n"
        for t in doing:
            report += f"  • {t}\n"
        report += "\n"
    if inbox:
        report += f"📥 Входящие ({len(inbox)}):\n"
        for t in inbox:
            report += f"  • {t}\n"
        report += "\n"

    if not any([done, review, doing, inbox]):
        report += "Задач на доске нет."

    return report
