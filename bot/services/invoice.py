import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "templates")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "output")


def _ensure_dirs():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def _set_cell_text(cell, text, bold=False, size=9, align=None):
    """Устанавливает текст ячейки с форматированием"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def _add_paragraph(doc, text, size=10, bold=False, align=None):
    """Добавляет форматированный параграф"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def _num_to_words(amount: float) -> str:
    """Простая конвертация числа в слова (тенге)"""
    # Базовые числа
    ones = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    ones_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
             "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]

    integer_part = int(amount)
    kopecks = round((amount - integer_part) * 100)

    if integer_part == 0:
        return f"Ноль тенге {kopecks:02d} тиын"

    parts = []

    # Миллионы
    if integer_part >= 1000000:
        m = integer_part // 1000000
        if m >= 100:
            parts.append(hundreds[m // 100])
            m = m % 100
        if 10 <= m <= 19:
            parts.append(teens[m - 10])
            m = 0
        elif m >= 20:
            parts.append(tens[m // 10])
            m = m % 10
        if m > 0:
            parts.append(ones[m])
        last = (integer_part // 1000000) % 10
        last2 = (integer_part // 1000000) % 100
        if last == 1 and last2 != 11:
            parts.append("миллион")
        elif 2 <= last <= 4 and not (12 <= last2 <= 14):
            parts.append("миллиона")
        else:
            parts.append("миллионов")
        integer_part = integer_part % 1000000

    # Тысячи
    if integer_part >= 1000:
        t = integer_part // 1000
        if t >= 100:
            parts.append(hundreds[t // 100])
            t = t % 100
        if 10 <= t <= 19:
            parts.append(teens[t - 10])
            t = 0
        elif t >= 20:
            parts.append(tens[t // 10])
            t = t % 10
        if t > 0:
            parts.append(ones_f[t])
        last = (integer_part // 1000) % 10
        last2 = (integer_part // 1000) % 100
        if last == 1 and last2 != 11:
            parts.append("тысяча")
        elif 2 <= last <= 4 and not (12 <= last2 <= 14):
            parts.append("тысячи")
        else:
            parts.append("тысяч")
        integer_part = integer_part % 1000

    # Сотни, десятки, единицы
    if integer_part >= 100:
        parts.append(hundreds[integer_part // 100])
        integer_part = integer_part % 100
    if 10 <= integer_part <= 19:
        parts.append(teens[integer_part - 10])
        integer_part = 0
    elif integer_part >= 20:
        parts.append(tens[integer_part // 10])
        integer_part = integer_part % 10
    if integer_part > 0:
        parts.append(ones[integer_part])

    result = " ".join(parts).strip()
    # Первая буква заглавная
    result = result[0].upper() + result[1:]
    return f"{result} тенге {kopecks:02d} тиын"


def generate_invoice(data: dict) -> str:
    """
    Генерирует счёт на оплату в формате ИП Knyaz.

    Обязательные поля в data:
    - НОМЕР_СЧЁТА: номер счёта
    - ДАТА_СЧЁТА: дата счёта
    - ПОКУПАТЕЛЬ_БИН: БИН/ИИН покупателя
    - ПОКУПАТЕЛЬ_НАЗВАНИЕ: название компании покупателя
    - ПОКУПАТЕЛЬ_АДРЕС: адрес покупателя
    - НОМЕР_ДОГОВОРА: номер договора
    - НАИМЕНОВАНИЕ_УСЛУГИ: название услуги
    - СУММА: сумма (число)
    """
    _ensure_dirs()
    doc = Document()

    # Устанавливаем узкие поля
    for section in doc.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # --- Предупреждение ---
    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warning.add_run(
        "Внимание! Оплата данного счета означает согласие с условиями поставки товара. "
        "Уведомление об оплате обязательно, в противном случае не гарантируется наличие "
        "товара на складе. Товар отпускается по факту прихода денег на р/с Поставщика, "
        "самовывозом, при наличии доверенности и документов удостоверяющих личность."
    )
    run.font.size = Pt(8)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.add_paragraph("")

    # --- Образец платежного поручения ---
    _add_paragraph(doc, "Образец платежного поручения", size=10, bold=True)

    # Таблица реквизитов поставщика
    table1 = doc.add_table(rows=3, cols=4)
    table1.style = "Table Grid"

    _set_cell_text(table1.rows[0].cells[0], "Бенефициар:", bold=True, size=9)
    _set_cell_text(table1.rows[1].cells[0], 'ИП "Knyaz"', size=9)
    _set_cell_text(table1.rows[2].cells[0], "ИИН: 050822551491", size=9)

    _set_cell_text(table1.rows[0].cells[2], "ИИК", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table1.rows[0].cells[3], "Кбе", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table1.rows[1].cells[2], "KZ53722S000031741265", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table1.rows[1].cells[3], "19", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    _set_cell_text(table1.rows[2].cells[0], "Банк бенефициара:", bold=True, size=9)
    # Объединяем ячейки для банка
    _set_cell_text(table1.rows[2].cells[2], "БИК", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table1.rows[2].cells[3], "Код назначения платежа", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Добавляем строку банка
    row_bank = table1.add_row()
    _set_cell_text(row_bank.cells[0], 'АО "KASPI BANK"', size=9)
    _set_cell_text(row_bank.cells[2], "CASPKZKA", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(row_bank.cells[3], "858", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("")

    # --- Заголовок счёта ---
    номер = data.get("НОМЕР_СЧЁТА", "—")
    дата = data.get("ДАТА_СЧЁТА", "—")
    _add_paragraph(doc, f"Счет на оплату № {номер} от {дата}", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Линия
    _add_paragraph(doc, "—" * 80, size=8)

    # --- Поставщик ---
    p = doc.add_paragraph()
    run_label = p.add_run("Поставщик:  ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.name = "Times New Roman"
    run_val = p.add_run(
        'БИН / ИИН 050822551491, ИП "Knyaz", 398 Казахстан Карагандинская обл. '
        'г. Караганда г. Караганда ул. улица Шахтеров д. 52б кв. (офис) 67'
    )
    run_val.font.size = Pt(10)
    run_val.font.name = "Times New Roman"

    # --- Покупатель ---
    покупатель_бин = data.get("ПОКУПАТЕЛЬ_БИН", "—")
    покупатель_название = data.get("ПОКУПАТЕЛЬ_НАЗВАНИЕ", "—")
    покупатель_адрес = data.get("ПОКУПАТЕЛЬ_АДРЕС", "—")

    p2 = doc.add_paragraph()
    run_label2 = p2.add_run("Покупатель:  ")
    run_label2.bold = True
    run_label2.font.size = Pt(10)
    run_label2.font.name = "Times New Roman"
    run_val2 = p2.add_run(
        f"БИН / ИИН {покупатель_бин}, {покупатель_название}, {покупатель_адрес}"
    )
    run_val2.font.size = Pt(10)
    run_val2.font.name = "Times New Roman"

    # --- Договор ---
    номер_договора = data.get("НОМЕР_ДОГОВОРА", "—")
    название_договора = data.get("НАЗВАНИЕ_ДОГОВОРА", "ДОГОВОР ВОЗМЕЗДНОГО ОКАЗАНИЯ УСЛУГ")

    p3 = doc.add_paragraph()
    run_label3 = p3.add_run("Договор:  ")
    run_label3.bold = True
    run_label3.font.size = Pt(10)
    run_label3.font.name = "Times New Roman"
    run_val3 = p3.add_run(f"{название_договора} №{номер_договора}")
    run_val3.font.size = Pt(10)
    run_val3.font.name = "Times New Roman"

    doc.add_paragraph("")

    # --- Таблица услуг ---
    наименование = data.get("НАИМЕНОВАНИЕ_УСЛУГИ", "—")
    сумма_str = data.get("СУММА", "0")
    # Очищаем сумму от текста
    сумма_clean = "".join(c for c in str(сумма_str) if c.isdigit() or c in ".,")
    сумма_clean = сумма_clean.replace(",", ".")
    try:
        сумма = float(сумма_clean)
    except ValueError:
        сумма = 0

    сумма_fmt = f"{сумма:,.2f}".replace(",", " ")
    количество = data.get("КОЛИЧЕСТВО", "1,000")
    единица = data.get("ЕДИНИЦА", "Одна услуга")
    код = data.get("КОД", "00000000009")

    table2 = doc.add_table(rows=2, cols=7)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    headers = ["№", "Код", "Наименование", "Кол-во", "Ед.", "Цена", "Сумма"]
    for i, h in enumerate(headers):
        _set_cell_text(table2.rows[0].cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Данные
    _set_cell_text(table2.rows[1].cells[0], "1", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table2.rows[1].cells[1], код, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table2.rows[1].cells[2], наименование, size=9)
    _set_cell_text(table2.rows[1].cells[3], количество, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table2.rows[1].cells[4], единица, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table2.rows[1].cells[5], сумма_fmt, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(table2.rows[1].cells[6], сумма_fmt, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Итого
    row_total = table2.add_row()
    # Объединяем первые 6 ячеек
    _set_cell_text(row_total.cells[5], "Итого:", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(row_total.cells[6], сумма_fmt, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("")

    # --- Итого текстом ---
    _add_paragraph(doc, f"Всего наименований 1, на сумму {сумма_fmt} KZT", size=10)

    # Сумма прописью
    сумма_прописью = _num_to_words(сумма)
    p_total = doc.add_paragraph()
    run_total_label = p_total.add_run("Всего к оплате: ")
    run_total_label.bold = True
    run_total_label.font.size = Pt(10)
    run_total_label.font.name = "Times New Roman"
    run_total_val = p_total.add_run(сумма_прописью)
    run_total_val.bold = True
    run_total_val.font.size = Pt(10)
    run_total_val.font.name = "Times New Roman"

    # Линия
    _add_paragraph(doc, "—" * 80, size=8)

    doc.add_paragraph("")

    # Подпись
    p_sign = doc.add_paragraph()
    run_sign = p_sign.add_run("Исполнитель _________________________________ /Директор/")
    run_sign.bold = True
    run_sign.font.size = Pt(10)
    run_sign.font.name = "Times New Roman"

    # Сохраняем
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"Счёт_{номер}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)

    logger.info(f"Счёт сгенерирован: {output_path}")
    return output_path
