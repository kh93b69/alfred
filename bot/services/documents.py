import os
import logging
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "templates")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "output")


def _ensure_dirs():
    """Создаёт папки если их нет"""
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_contract_template():
    """Создаёт шаблон договора оказания услуг"""
    _ensure_dirs()
    doc = Document()

    # Стиль
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ДОГОВОР ОКАЗАНИЯ УСЛУГ №{{НОМЕР_ДОГОВОРА}}")
    run.bold = True
    run.font.size = Pt(14)

    # Дата и город
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("г. {{ГОРОД}}")
    p.add_run("\t\t\t\t\t\t{{ДАТА_ДОГОВОРА}}")

    doc.add_paragraph("")

    # Преамбула
    doc.add_paragraph(
        '{{НАЗВАНИЕ_ИСПОЛНИТЕЛЯ}}, в лице {{ДОЛЖНОСТЬ_ИСПОЛНИТЕЛЯ}} {{ФИО_ИСПОЛНИТЕЛЯ}}, '
        'действующего на основании {{ОСНОВАНИЕ_ИСПОЛНИТЕЛЯ}}, именуемое в дальнейшем '
        '«Исполнитель», с одной стороны, и '
        '{{НАЗВАНИЕ_ЗАКАЗЧИКА}}, в лице {{ДОЛЖНОСТЬ_ЗАКАЗЧИКА}} {{ФИО_ЗАКАЗЧИКА}}, '
        'действующего на основании {{ОСНОВАНИЕ_ЗАКАЗЧИКА}}, именуемое в дальнейшем '
        '«Заказчик», с другой стороны, заключили настоящий договор о нижеследующем:'
    )

    # Раздел 1
    h1 = doc.add_paragraph()
    run = h1.add_run("1. ПРЕДМЕТ ДОГОВОРА")
    run.bold = True

    doc.add_paragraph(
        "1.1. Исполнитель обязуется оказать Заказчику следующие услуги: {{ОПИСАНИЕ_УСЛУГ}}"
    )
    doc.add_paragraph(
        "1.2. Срок оказания услуг: с {{ДАТА_НАЧАЛА}} по {{ДАТА_ОКОНЧАНИЯ}}."
    )

    # Раздел 2
    h2 = doc.add_paragraph()
    run = h2.add_run("2. СТОИМОСТЬ УСЛУГ И ПОРЯДОК РАСЧЁТОВ")
    run.bold = True

    doc.add_paragraph(
        "2.1. Стоимость услуг по настоящему договору составляет {{СУММА}} ({{СУММА_ПРОПИСЬЮ}}) рублей."
    )
    doc.add_paragraph(
        "2.2. Оплата производится в течение {{СРОК_ОПЛАТЫ}} банковских дней с момента подписания акта оказанных услуг."
    )

    # Раздел 3
    h3 = doc.add_paragraph()
    run = h3.add_run("3. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    run.bold = True

    doc.add_paragraph("Исполнитель:")
    doc.add_paragraph("{{НАЗВАНИЕ_ИСПОЛНИТЕЛЯ}}")
    doc.add_paragraph("ИНН: {{ИНН_ИСПОЛНИТЕЛЯ}}")
    doc.add_paragraph("Р/с: {{РАСЧЁТНЫЙ_СЧЁТ_ИСПОЛНИТЕЛЯ}}")
    doc.add_paragraph("Банк: {{БАНК_ИСПОЛНИТЕЛЯ}}")
    doc.add_paragraph("БИК: {{БИК_ИСПОЛНИТЕЛЯ}}")

    doc.add_paragraph("")
    doc.add_paragraph("Заказчик:")
    doc.add_paragraph("{{НАЗВАНИЕ_ЗАКАЗЧИКА}}")
    doc.add_paragraph("ИНН: {{ИНН_ЗАКАЗЧИКА}}")
    doc.add_paragraph("Р/с: {{РАСЧЁТНЫЙ_СЧЁТ_ЗАКАЗЧИКА}}")
    doc.add_paragraph("Банк: {{БАНК_ЗАКАЗЧИКА}}")
    doc.add_paragraph("БИК: {{БИК_ЗАКАЗЧИКА}}")

    filepath = os.path.join(TEMPLATES_DIR, "договор_услуги.docx")
    doc.save(filepath)
    logger.info(f"Создан шаблон договора: {filepath}")
    return filepath


def create_invoice_template():
    """Создаёт шаблон счёта на оплату"""
    _ensure_dirs()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СЧЁТ НА ОПЛАТУ №{{НОМЕР_СЧЁТА}} от {{ДАТА_СЧЁТА}}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    # Поставщик
    doc.add_paragraph("Поставщик: {{НАЗВАНИЕ_ПОСТАВЩИКА}}")
    doc.add_paragraph("ИНН: {{ИНН_ПОСТАВЩИКА}} / КПП: {{КПП_ПОСТАВЩИКА}}")
    doc.add_paragraph("Р/с: {{РАСЧЁТНЫЙ_СЧЁТ_ПОСТАВЩИКА}}")
    doc.add_paragraph("Банк: {{БАНК_ПОСТАВЩИКА}}, БИК: {{БИК_ПОСТАВЩИКА}}")
    doc.add_paragraph("К/с: {{КОРР_СЧЁТ_ПОСТАВЩИКА}}")

    doc.add_paragraph("")

    # Покупатель
    doc.add_paragraph("Покупатель: {{НАЗВАНИЕ_ПОКУПАТЕЛЯ}}")
    doc.add_paragraph("ИНН: {{ИНН_ПОКУПАТЕЛЯ}} / КПП: {{КПП_ПОКУПАТЕЛЯ}}")
    doc.add_paragraph("Адрес: {{АДРЕС_ПОКУПАТЕЛЯ}}")

    doc.add_paragraph("")

    # Таблица
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"

    # Заголовки таблицы
    headers = ["№", "Наименование", "Кол-во", "Цена", "Сумма"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Строка с данными
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "{{НАИМЕНОВАНИЕ_ТОВАРА}}"
    table.rows[1].cells[2].text = "{{КОЛИЧЕСТВО}}"
    table.rows[1].cells[3].text = "{{ЦЕНА}}"
    table.rows[1].cells[4].text = "{{СУММА}}"

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Итого: {{СУММА}} руб.")
    run.bold = True

    p2 = doc.add_paragraph()
    run2 = p2.add_run("В том числе НДС: {{НДС}} руб.")

    doc.add_paragraph("")
    doc.add_paragraph("Руководитель _________________ / {{ФИО_РУКОВОДИТЕЛЯ}} /")
    doc.add_paragraph("")
    doc.add_paragraph("Бухгалтер _________________ / {{ФИО_БУХГАЛТЕРА}} /")

    filepath = os.path.join(TEMPLATES_DIR, "счёт.docx")
    doc.save(filepath)
    logger.info(f"Создан шаблон счёта: {filepath}")
    return filepath


def list_templates() -> list[str]:
    """Возвращает список доступных шаблонов"""
    _ensure_dirs()
    return [f for f in os.listdir(TEMPLATES_DIR) if f.endswith(".docx")]


def get_template_fields(template_name: str) -> list[str]:
    """Извлекает все плейсхолдеры из шаблона"""
    filepath = os.path.join(TEMPLATES_DIR, template_name)
    doc = Document(filepath)

    fields = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Ищем {{ПОЛЕ}}
        import re
        found = re.findall(r"\{\{(.+?)\}\}", text)
        fields.update(found)

    # Проверяем таблицы тоже
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                import re
                found = re.findall(r"\{\{(.+?)\}\}", cell.text)
                fields.update(found)

    return sorted(fields)


def fill_template(template_name: str, data: dict[str, str]) -> str:
    """
    Заполняет шаблон данными и сохраняет результат.
    Возвращает путь к готовому файлу.
    """
    _ensure_dirs()
    filepath = os.path.join(TEMPLATES_DIR, template_name)
    doc = Document(filepath)

    # Заменяем плейсхолдеры в параграфах
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                # Заменяем в каждом run чтобы сохранить форматирование
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                # Если placeholder разбит по runs — заменяем в полном тексте
                if placeholder in paragraph.text:
                    full = paragraph.text
                    full = full.replace(placeholder, value)
                    # Очищаем и перезаписываем
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = full

    # Заменяем в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)

    # Сохраняем результат
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    name = template_name.rsplit(".", 1)[0]
    output_name = f"{name}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)

    logger.info(f"Документ сгенерирован: {output_path}")
    return output_path


# Создаём шаблоны при первом импорте если их нет
_ensure_dirs()
if not list_templates():
    create_contract_template()
    create_invoice_template()
